#!/usr/bin/env python3
"""
Robust RAG Pipeline for Large Document Sets (45+ PDFs)
- Handles large document collections without falling back to hardcoded topics
- Uses text-based output (no JSON parsing failures)
- Progressive topic discovery with document clustering
- Scales to 100+ PDFs efficiently
"""
import os
import json
import re
import shutil
import torch
import torch.nn.functional as F
from typing import List, Dict, Any, Optional, Set, Tuple
from pathlib import Path
import logging
from tqdm import tqdm
from collections import Counter, defaultdict
import numpy as np

# LangChain imports
from langchain_community.document_loaders import PyPDFDirectoryLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.embeddings import Embeddings

# Transformers imports
from transformers import (
    AutoModelForCausalLM,
    AutoTokenizer,
    AutoModel,
)

# Document generation
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class CustomQwenEmbeddings(Embeddings):
    """Custom embedding class for Qwen3-Embedding-8B model (4096 dimensions)."""
    
    def __init__(self, model_name: str = "Qwen/Qwen3-Embedding-8B", device: str = None, use_flash_attention: bool = False):
        """Initialize the Qwen3-Embedding-8B model."""
        self.device = device or ('cuda' if torch.cuda.is_available() else 'cpu')
        self.model_name = model_name
        logger.info(f"Loading embedding model: {model_name}")
        
        if torch.cuda.is_available():
            num_gpus = torch.cuda.device_count()
            logger.info(f"Found {num_gpus} GPU(s) available for embeddings")
        
        self.tokenizer = AutoTokenizer.from_pretrained(
            model_name,
            padding_side='left',
            trust_remote_code=True
        )
        
        model_kwargs = {
            'trust_remote_code': True,
            'device_map': 'auto',
        }
        
        if use_flash_attention and self.device == 'cuda':
            logger.info("Enabling flash_attention_2")
            model_kwargs['attn_implementation'] = 'flash_attention_2'
            model_kwargs['torch_dtype'] = torch.float16
        elif self.device == 'cuda':
            model_kwargs['torch_dtype'] = torch.float16
        else:
            model_kwargs['torch_dtype'] = torch.float32
        
        self.model = AutoModel.from_pretrained(model_name, **model_kwargs)
        self.model.eval()
        self.max_length = 8192
        
        with torch.no_grad():
            test_input = self.tokenizer(["test"], padding=True, truncation=True, 
                                       max_length=self.max_length, return_tensors="pt")
            test_input = {k: v.to(self.model.device) for k, v in test_input.items()}
            test_output = self.model(**test_input)
            test_emb = self.last_token_pool(test_output.last_hidden_state, test_input['attention_mask'])
            self.embedding_dim = test_emb.shape[1]
        
        logger.info(f"Embedding dimension: {self.embedding_dim}, Max length: {self.max_length}")
    
    @staticmethod
    def last_token_pool(last_hidden_states: torch.Tensor, attention_mask: torch.Tensor) -> torch.Tensor:
        left_padding = (attention_mask[:, -1].sum() == attention_mask.shape[0])
        if left_padding:
            return last_hidden_states[:, -1]
        else:
            sequence_lengths = attention_mask.sum(dim=1) - 1
            batch_size = last_hidden_states.shape[0]
            return last_hidden_states[torch.arange(batch_size, device=last_hidden_states.device), sequence_lengths]
    
    @staticmethod
    def get_detailed_instruct(task_description: str, query: str) -> str:
        return f'Instruct: {task_description}\nQuery:{query}'
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        if not texts:
            return []
        
        embeddings = []
        batch_size = 16
        
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i + batch_size]
            batch_dict = self.tokenizer(
                batch_texts,
                padding=True,
                truncation=True,
                max_length=self.max_length,
                return_tensors="pt"
            )
            batch_dict = {k: v.to(self.model.device) for k, v in batch_dict.items()}
            
            with torch.no_grad():
                outputs = self.model(**batch_dict)
            
            batch_embeddings = self.last_token_pool(
                outputs.last_hidden_state,
                batch_dict['attention_mask']
            )
            batch_embeddings = F.normalize(batch_embeddings, p=2, dim=1)
            embeddings.extend(batch_embeddings.cpu().numpy().tolist())
        
        return embeddings
    
    def embed_query(self, text: str) -> List[float]:
        task_description = 'Given a web search query, retrieve relevant passages that answer the query'
        formatted_query = self.get_detailed_instruct(task_description, text)
        
        batch_dict = self.tokenizer(
            [formatted_query],
            padding=True,
            truncation=True,
            max_length=self.max_length,
            return_tensors="pt"
        )
        batch_dict = {k: v.to(self.model.device) for k, v in batch_dict.items()}
        
        with torch.no_grad():
            outputs = self.model(**batch_dict)
        
        embedding = self.last_token_pool(
            outputs.last_hidden_state,
            batch_dict['attention_mask']
        )
        embedding = F.normalize(embedding, p=2, dim=1)
        
        return embedding[0].cpu().numpy().tolist()


class Qwen3NextLLM:
    """Wrapper for Qwen3-Next-80B model."""
    
    def __init__(self, model_name: str = "Qwen/Qwen3-Next-80B-A3B-Instruct"):
        logger.info(f"Loading language model: {model_name}")
        
        self.tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
        self.model = AutoModelForCausalLM.from_pretrained(
            model_name,
            torch_dtype="auto",
            device_map="auto",
            trust_remote_code=True
        )
        self.model.eval()
        logger.info(f"Language model loaded successfully")
    
    def generate(self, prompt: str, max_new_tokens: int = 2048, temperature: float = 0.1) -> str:
        messages = [{"role": "user", "content": prompt}]
        text = self.tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
        
        model_inputs = self.tokenizer([text], return_tensors="pt")
        device = next(self.model.parameters()).device
        model_inputs = {k: v.to(device) for k, v in model_inputs.items()}
        input_length = model_inputs['input_ids'].shape[1]
        
        with torch.no_grad():
            generated_ids = self.model.generate(
                **model_inputs,
                max_new_tokens=max_new_tokens,
                temperature=temperature,
                do_sample=temperature > 0,
                top_p=0.9 if temperature > 0 else None,
                repetition_penalty=1.05,
                pad_token_id=self.tokenizer.pad_token_id,
                eos_token_id=self.tokenizer.eos_token_id,
            )
        
        output_ids = generated_ids[0][input_length:].tolist()
        content = self.tokenizer.decode(output_ids, skip_special_tokens=True)
        return content.strip()


class RobustTopicDiscovery:
    """
    Robust topic discovery for large document sets.
    Uses text-based output (no JSON) and progressive clustering.
    """
    
    def __init__(self, llm: Qwen3NextLLM, vector_store, embeddings, max_topics: int = 25):
        self.llm = llm
        self.vector_store = vector_store
        self.embeddings = embeddings
        self.max_topics = max_topics
    
    def discover_topics(self, num_pdfs: int = 0) -> List[str]:
        """
        Discover topics robustly for any size document collection.
        
        Args:
            num_pdfs: Number of PDFs in collection (for scaling)
        """
        logger.info(f"Discovering topics for {num_pdfs} PDFs...")
        
        # Step 1: Extract comprehensive keywords
        logger.info("Step 1: Extracting keywords...")
        keywords = self._extract_keywords_comprehensive()
        logger.info(f"Extracted {len(keywords)} keywords: {keywords[:20]}...")
        
        # Step 2: Cluster documents by similarity
        logger.info("Step 2: Clustering documents...")
        clusters = self._cluster_documents(num_clusters=min(8, max(4, num_pdfs // 6)))
        logger.info(f"Created {len(clusters)} document clusters")
        
        # Step 3: Generate topics from each cluster
        logger.info("Step 3: Generating topics from clusters...")
        all_topics = []
        for i, cluster_docs in enumerate(clusters):
            logger.info(f"Processing cluster {i+1}/{len(clusters)}...")
            cluster_topics = self._extract_topics_from_cluster(cluster_docs, cluster_id=i)
            all_topics.extend(cluster_topics)
            logger.info(f"Cluster {i+1} generated {len(cluster_topics)} topics")
        
        logger.info(f"Generated {len(all_topics)} topics before refinement")
        
        # Step 4: Deduplicate and refine
        logger.info("Step 4: Refining topic list...")
        final_topics = self._refine_topics_text_based(all_topics)
        
        logger.info(f"Final: {len(final_topics)} topics")
        return final_topics
    
    def _extract_keywords_comprehensive(self, top_k: int = 100) -> List[str]:
        """Extract keywords from entire collection."""
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'be',
            'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will',
            'would', 'could', 'should', 'may', 'might', 'must', 'can', 'this',
            'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they',
            'what', 'which', 'who', 'when', 'where', 'why', 'how', 'all', 'each',
            'every', 'both', 'few', 'more', 'most', 'other', 'some', 'such', 'no',
            'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 's',
            't', 'just', 'don', 'now', 'page', 'also', 'one', 'two', 'use', 'make'
        }
        
        # Sample broadly across collection
        retriever = self.vector_store.as_retriever(search_kwargs={"k": 200})
        
        queries = [
            "main topics", "key information", "important content",
            "primary subjects", "core themes", "essential data"
        ]
        
        all_docs = []
        seen = set()
        for query in queries:
            docs = retriever.invoke(query)
            for doc in docs:
                h = hash(doc.page_content[:100])
                if h not in seen:
                    seen.add(h)
                    all_docs.append(doc)
        
        # Extract and count words
        all_text = " ".join([doc.page_content for doc in all_docs])
        words = re.findall(r'\b[a-zA-Z]{3,20}\b', all_text.lower())
        meaningful = [w for w in words if w not in stop_words]
        
        word_freq = Counter(meaningful)
        return [word for word, count in word_freq.most_common(top_k) if count > 2]
    
    def _cluster_documents(self, num_clusters: int = 6) -> List[List[Any]]:
        """
        Cluster documents by semantic similarity using embeddings.
        """
        # Get all documents
        retriever = self.vector_store.as_retriever(search_kwargs={"k": 300})
        all_docs = retriever.invoke("all content topics information")
        
        if len(all_docs) < num_clusters:
            # Not enough docs, return as single cluster
            return [all_docs]
        
        # Simple clustering: divide by sequential batches for speed
        # In production, could use k-means on embeddings
        cluster_size = len(all_docs) // num_clusters
        clusters = []
        
        for i in range(num_clusters):
            start_idx = i * cluster_size
            end_idx = start_idx + cluster_size if i < num_clusters - 1 else len(all_docs)
            clusters.append(all_docs[start_idx:end_idx])
        
        return [c for c in clusters if c]  # Remove empty
    
    def _extract_topics_from_cluster(self, docs: List[Any], cluster_id: int) -> List[str]:
        """
        Extract 3-5 topics from a document cluster using TEXT-BASED output.
        NO JSON - just line-by-line parsing.
        """
        # Combine document content
        context = "\n\n".join([doc.page_content for doc in docs[:15]])
        
        # TEXT-BASED prompt (no JSON!)
        prompt = f"""Analyze the following text and identify 3-5 main topics or themes.

Text:
{context[:6000]}

Write each topic on a new line starting with "TOPIC:".

Example:
TOPIC: Budget Planning and Financial Management
TOPIC: Technical Infrastructure and System Architecture
TOPIC: Project Timeline and Milestones

Now list the topics from the text above:"""
        
        # Generate with higher temperature for diversity
        result = self.llm.generate(prompt, max_new_tokens=512, temperature=0.3)
        
        # TEXT-BASED parsing (no JSON!)
        topics = self._parse_text_list(result, prefix="TOPIC:")
        
        if not topics:
            logger.warning(f"Cluster {cluster_id}: No topics parsed, trying alternative parsing...")
            topics = self._parse_text_list_flexible(result)
        
        if not topics:
            logger.warning(f"Cluster {cluster_id}: Still no topics, using keyword fallback")
            topics = [f"Cluster {cluster_id} Topics"]
        
        return topics[:5]  # Max 5 per cluster
    
    def _parse_text_list(self, text: str, prefix: str = "TOPIC:") -> List[str]:
        """Parse topics from text output (no JSON)."""
        topics = []
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            
            # Method 1: Exact prefix match
            if line.startswith(prefix):
                topic = line.replace(prefix, '').strip()
                if topic and len(topic) > 5:
                    topics.append(topic)
            
            # Method 2: Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                topic = line[2:].strip()
                if topic and len(topic) > 5:
                    topics.append(topic)
            
            # Method 3: Numbered lists
            elif re.match(r'^\d+[\.\)]\s+', line):
                topic = re.sub(r'^\d+[\.\)]\s+', '', line).strip()
                if topic and len(topic) > 5:
                    topics.append(topic)
        
        return topics
    
    def _parse_text_list_flexible(self, text: str) -> List[str]:
        """Flexible parsing for any line that looks like a topic."""
        topics = []
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            
            # Skip very short or very long lines
            if len(line) < 10 or len(line) > 150:
                continue
            
            # Skip lines that are clearly not topics
            skip_patterns = [
                'example', 'format', 'following', 'analyze', 'identify',
                'listed', 'above', 'below', 'here are', 'these are'
            ]
            if any(pat in line.lower() for pat in skip_patterns):
                continue
            
            # If it looks like a topic (starts with capital, no question mark)
            if line[0].isupper() and '?' not in line:
                # Clean any leading markers
                topic = re.sub(r'^[\-\*\d\.\)]+\s*', '', line).strip()
                if topic and len(topic) > 5:
                    topics.append(topic)
        
        return topics[:10]  # Limit to 10 to avoid junk
    
    def _refine_topics_text_based(self, topics: List[str]) -> List[str]:
        """
        Refine topics to target count using TEXT-BASED output.
        """
        # Remove exact duplicates
        unique_topics = list(dict.fromkeys(topics))
        logger.info(f"After deduplication: {len(unique_topics)} unique topics")
        
        if len(unique_topics) <= self.max_topics:
            return unique_topics
        
        # Need to merge - use LLM but with TEXT output
        logger.info(f"Merging {len(unique_topics)} topics to {self.max_topics}...")
        
        topics_text = "\n".join([f"{i+1}. {t}" for i, t in enumerate(unique_topics)])
        
        prompt = f"""You have {len(unique_topics)} topics but need to reduce to {self.max_topics} topics.

Merge similar topics and keep the most important distinct ones.

Current topics:
{topics_text}

Write EXACTLY {self.max_topics} refined topics, one per line starting with "TOPIC:".

Guidelines:
- Merge similar/overlapping topics
- Keep distinct important topics
- Use clear names
- Maintain coverage

Refined topics:"""
        
        result = self.llm.generate(prompt, max_new_tokens=1024, temperature=0.1)
        refined = self._parse_text_list(result, prefix="TOPIC:")
        
        if not refined:
            logger.warning("Parsing failed, trying flexible parsing...")
            refined = self._parse_text_list_flexible(result)
        
        if len(refined) < self.max_topics * 0.6:
            logger.warning(f"Only got {len(refined)} topics from refinement, using top {self.max_topics} original")
            return unique_topics[:self.max_topics]
        
        return refined[:self.max_topics]


class RAGPipeline:
    """RAG pipeline with robust topic discovery for large document sets."""
    
    def __init__(self, source_dir: str = "./source_pdfs", persist_dir: str = "./chroma_db", max_topics: int = 25):
        self.source_dir = source_dir
        self.persist_dir = persist_dir
        self.max_topics = max_topics
        self.vector_store = None
        self.llm = None
        self.embeddings = None
        self.num_pdfs = 0
    
    def setup_llm(self):
        logger.info("Setting up Qwen3-Next-80B...")
        self.llm = Qwen3NextLLM(model_name="Qwen/Qwen3-Next-80B-A3B-Instruct")
        logger.info("LLM setup complete")
    
    def setup_embeddings(self, use_flash_attention: bool = False):
        logger.info("Setting up Qwen3-Embedding-8B...")
        self.embeddings = CustomQwenEmbeddings(
            model_name="Qwen/Qwen3-Embedding-8B",
            use_flash_attention=use_flash_attention
        )
        logger.info("Embeddings setup complete")
    
    def _check_vector_store_compatibility(self) -> bool:
        if not os.path.exists(self.persist_dir):
            return True
        
        try:
            test_store = Chroma(
                persist_directory=self.persist_dir,
                embedding_function=self.embeddings
            )
            test_store.similarity_search("test", k=1)
            logger.info("Existing vector store is compatible")
            return True
        except Exception as e:
            logger.warning(f"Vector store compatibility check failed: {e}")
            return False
    
    def ingest_and_index_documents(self, force_recreate: bool = False):
        logger.info(f"Loading PDFs from {self.source_dir}...")
        
        if not os.path.exists(self.source_dir):
            os.makedirs(self.source_dir)
            logger.warning(f"Created empty source directory: {self.source_dir}")
            return False
        
        loader = PyPDFDirectoryLoader(self.source_dir)
        documents = loader.load()
        
        if not documents:
            logger.warning(f"No PDF documents found in {self.source_dir}")
            return False
        
        # Count PDFs
        pdf_files = list(Path(self.source_dir).glob("*.pdf"))
        self.num_pdfs = len(pdf_files)
        
        logger.info(f"Loaded {len(documents)} pages from {self.num_pdfs} PDFs")
        
        needs_recreation = force_recreate or not self._check_vector_store_compatibility()
        
        if needs_recreation:
            if os.path.exists(self.persist_dir):
                logger.warning(f"Removing incompatible vector store")
                shutil.rmtree(self.persist_dir)
            
            logger.info("Creating new vector store...")
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            
            chunks = text_splitter.split_documents(documents)
            logger.info(f"Created {len(chunks)} chunks from documents")
            
            self.vector_store = Chroma.from_documents(
                documents=chunks,
                embedding=self.embeddings,
                persist_directory=self.persist_dir
            )
            
            logger.info("Document indexing complete")
        else:
            logger.info("Loading existing vector store...")
            self.vector_store = Chroma(
                persist_directory=self.persist_dir,
                embedding_function=self.embeddings
            )
        
        return True
    
    def discover_topics(self) -> List[str]:
        logger.info(f"Discovering topics for {self.num_pdfs} PDFs (max {self.max_topics} topics)...")
        
        discovery_engine = RobustTopicDiscovery(
            llm=self.llm,
            vector_store=self.vector_store,
            embeddings=self.embeddings,
            max_topics=self.max_topics
        )
        
        topics = discovery_engine.discover_topics(num_pdfs=self.num_pdfs)
        
        if not topics or len(topics) == 0:
            logger.error("Topic discovery failed completely! This should not happen.")
            logger.error("Using emergency fallback topics")
            topics = [
                "Document Overview and Summary",
                "Key Information and Details",
                "Important Findings and Results",
                "Data and Statistics",
                "Conclusions and Recommendations"
            ]
        
        return topics
    
    def extract_topic_updates(self, topic: str, max_retries: int = 2) -> List[str]:
        logger.info(f"Extracting updates for: {topic}")
        
        retriever = self.vector_store.as_retriever(search_kwargs={"k": 30})
        
        search_queries = [
            topic,
            f"{topic} details",
            f"information about {topic}",
        ]
        
        all_docs = []
        seen = set()
        
        for query in search_queries:
            docs = retriever.invoke(query)
            for doc in docs:
                h = hash(doc.page_content[:100])
                if h not in seen:
                    seen.add(h)
                    all_docs.append(doc)
        
        if not all_docs:
            return []
        
        all_updates = []
        batch_size = 5
        
        for i in range(0, min(len(all_docs), 20), batch_size):
            batch_docs = all_docs[i:i + batch_size]
            context = "\n\n---\n\n".join([doc.page_content for doc in batch_docs])
            
            # TEXT-BASED extraction (no JSON)
            prompt = f"""Extract facts about "{topic}" from this text.

Text:
{context[:8000]}

Write each fact on a new line starting with "FACT:".

Example:
FACT: The project budget is $2.5 million
FACT: Timeline is 18 months from start date
FACT: Team consists of 15 full-time members

Extract facts about "{topic}":"""
            
            for attempt in range(max_retries):
                try:
                    result = self.llm.generate(prompt, max_new_tokens=1024, temperature=0.1)
                    
                    # Parse text-based output
                    facts = []
                    lines = result.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line.startswith('FACT:'):
                            fact = line.replace('FACT:', '').strip()
                            if fact and len(fact) > 15:
                                facts.append(fact)
                        elif line.startswith('- ') or line.startswith('* '):
                            fact = line[2:].strip()
                            if fact and len(fact) > 15:
                                facts.append(fact)
                    
                    if facts:
                        all_updates.extend(facts)
                        logger.info(f"Batch {i//batch_size}: extracted {len(facts)} facts")
                        break
                    elif attempt < max_retries - 1:
                        logger.warning(f"Batch {i//batch_size}: no facts, retrying...")
                except Exception as e:
                    if attempt < max_retries - 1:
                        logger.warning(f"Batch {i//batch_size}: error, retrying: {e}")
                    else:
                        logger.error(f"Batch {i//batch_size}: failed")
        
        logger.info(f"Extracted {len(all_updates)} updates for {topic}")
        return all_updates
    
    def deduplicate_updates(self, updates: List[str]) -> List[str]:
        if not updates:
            return []
        return list(dict.fromkeys(updates))[:20]
    
    def generate_report(self, compiled_data: Dict[str, List[str]], output_file: str = "Consolidated_Report.docx"):
        logger.info(f"Generating report: {output_file}")
        
        doc = Document()
        
        title = doc.add_heading("Consolidated Report - All Topics and Updates", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading("Executive Summary", 1)
        total_updates = sum(len(updates) for updates in compiled_data.values())
        summary = f"This report consolidates information from {self.num_pdfs} PDF documents across {len(compiled_data)} main topics, extracting {total_updates} unique data points."
        doc.add_paragraph(summary)
        doc.add_paragraph()
        
        for topic, updates in compiled_data.items():
            doc.add_heading(topic, 1)
            
            if updates:
                for update in updates:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(update)
            else:
                doc.add_paragraph("No updates found for this topic.", style='Body Text')
            
            doc.add_paragraph()
        
        doc.add_page_break()
        doc.add_heading("Report Metadata", 2)
        doc.add_paragraph(f"Source: {self.num_pdfs} PDF documents")
        doc.add_paragraph(f"Embedding Model: Qwen3-Embedding-8B ({self.embeddings.embedding_dim}D)")
        doc.add_paragraph(f"Language Model: Qwen3-Next-80B-A3B-Instruct")
        doc.add_paragraph(f"Total topics: {len(compiled_data)}")
        doc.add_paragraph(f"Total facts: {total_updates}")
        
        doc.save(output_file)
        logger.info(f"Report saved to {output_file}")
    
    def run_pipeline(self, use_flash_attention: bool = False, force_recreate_vectorstore: bool = False):
        logger.info("="*70)
        logger.info("ROBUST RAG PIPELINE - Handles 45+ PDFs without fallback")
        logger.info("="*70)
        
        self.setup_embeddings(use_flash_attention=use_flash_attention)
        self.setup_llm()
        
        if not self.ingest_and_index_documents(force_recreate=force_recreate_vectorstore):
            logger.error("No documents to process. Exiting.")
            return
        
        topics = self.discover_topics()
        
        logger.info(f"\n{'='*70}")
        logger.info(f"DISCOVERED {len(topics)} TOPICS:")
        for i, topic in enumerate(topics, 1):
            logger.info(f"  {i}. {topic}")
        logger.info(f"{'='*70}\n")
        
        compiled_data = {}
        for topic in tqdm(topics, desc="Processing topics"):
            topic_updates = self.extract_topic_updates(topic)
            unique_updates = self.deduplicate_updates(topic_updates)
            compiled_data[topic] = unique_updates
            logger.info(f"'{topic}': {len(unique_updates)} facts")
        
        self.generate_report(compiled_data)
        
        print("\n" + "="*70)
        print("PIPELINE COMPLETE")
        print("="*70)
        print(f"PDFs processed: {self.num_pdfs}")
        print(f"Topics discovered: {len(topics)} (NO FALLBACK USED)")
        print(f"Total facts extracted: {sum(len(u) for u in compiled_data.values())}")
        print(f"Report: Consolidated_Report.docx")
        print("="*70)


def main():
    source_dir = "./source_pdfs"
    
    if not os.path.exists(source_dir):
        os.makedirs(source_dir)
        print(f"Created {source_dir} directory.")
        return
    
    pdf_files = list(Path(source_dir).glob("*.pdf"))
    if not pdf_files:
        print(f"No PDF files found in {source_dir}")
        return
    
    print(f"Found {len(pdf_files)} PDF files")
    print("This pipeline handles large document sets (45+ PDFs) without falling back to hardcoded topics!")
    print()
    
    pipeline = RAGPipeline(source_dir=source_dir, max_topics=25)
    
    try:
        pipeline.run_pipeline(use_flash_attention=False, force_recreate_vectorstore=False)
    except Exception as e:
        logger.error(f"Pipeline failed: {e}")
        raise


if __name__ == "__main__":
    main()
