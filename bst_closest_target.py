import pandas as pd
import numpy as np
import re
import os
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
from docx import Document
from docx.shared import Inches
from collections import Counter, defaultdict
import json
from itertools import combinations
import warnings
import gc  # For garbage collection
warnings.filterwarnings('ignore')

class LargeFileAnalyzer:
    def __init__(self, csv_path, chunk_size=10000, max_rows=None):
        self.csv_path = csv_path
        self.chunk_size = chunk_size
        self.max_rows = max_rows
        self.stats = defaultdict(lambda: defaultdict(dict))
        self.app_counters = {
            'base': Counter(),
            'variant': Counter(),
            'base_combos': Counter(),
            'variant_combos': Counter()
        }
        
    def filter_row(self, row):
        """Filter out rows based on conditions"""
        if row.get('release_type') == 'SMR' or row.get('vr_qualified') == 'TRUE':
            return False
        return True
    
    def clean_app_columns(self, row):
        """Clean app columns - replace [] with empty"""
        for col in ['priv_app_in_base_only', 'priv_app_in_variant_only']:
            if col in row and row[col] in ['[]', '[ ]']:
                row[col] = ''
        return row
    
    def extract_fingerprint(self, text):
        """Extract items between first and second /"""
        if pd.isna(text) or text == '':
            return text
        text = str(text)
        parts = text.split('/')
        if len(parts) >= 2:
            return parts[1]
        return text
    
    def parse_app_list(self, text):
        """Parse app list string into Python list"""
        if pd.isna(text) or text == '' or text == '[]':
            return []
        
        try:
            if isinstance(text, str):
                text = text.strip()
                if text.startswith('[') and text.endswith(']'):
                    text_json = text.replace("'", '"')
                    try:
                        return json.loads(text_json)
                    except:
                        return eval(text)
                else:
                    return [item.strip() for item in text.split(',') if item.strip()]
        except:
            return []
        
    def process_chunk(self, chunk):
        """Process a single chunk of data"""
        # Apply filtering
        chunk = chunk[chunk.apply(self.filter_row, axis=1)]
        
        # Clean app columns
        chunk = chunk.apply(self.clean_app_columns, axis=1)
        
        # Extract fingerprints
        chunk['fingerprint_extracted'] = chunk['fingerprint'].apply(self.extract_fingerprint)
        chunk['same_device_fingerprint_extracted'] = chunk['same_device_fingerprint'].apply(self.extract_fingerprint)
        
        # Parse app lists
        chunk['base_apps'] = chunk['priv_app_in_base_only'].apply(self.parse_app_list)
        chunk['variant_apps'] = chunk['priv_app_in_variant_only'].apply(self.parse_app_list)
        
        return chunk
    
    def update_statistics(self, chunk):
        """Update running statistics with chunk data"""
        # Group by device and release_type
        for (device, release_type), group in chunk.groupby(['device', 'release_type']):
            if device not in self.stats:
                self.stats[device] = {}
            if release_type not in self.stats[device]:
                self.stats[device][release_type] = {
                    'count': 0,
                    'fingerprint_stats': defaultdict(lambda: {'base': Counter(), 'variant': Counter()}),
                    'same_fingerprint_stats': defaultdict(lambda: {'base': Counter(), 'variant': Counter()})
                }
            
            self.stats[device][release_type]['count'] += len(group)
            
            # Process each row in the group
            for _, row in group.iterrows():
                # Update fingerprint stats
                fp = row.get('fingerprint_extracted', '')
                same_fp = row.get('same_device_fingerprint_extracted', '')
                
                # Count apps by fingerprint
                if fp:
                    for app in row['base_apps']:
                        self.stats[device][release_type]['fingerprint_stats'][fp]['base'][app] += 1
                    for app in row['variant_apps']:
                        self.stats[device][release_type]['fingerprint_stats'][fp]['variant'][app] += 1
                
                if same_fp:
                    for app in row['base_apps']:
                        self.stats[device][release_type]['same_fingerprint_stats'][same_fp]['base'][app] += 1
                    for app in row['variant_apps']:
                        self.stats[device][release_type]['same_fingerprint_stats'][same_fp]['variant'][app] += 1
                
                # Update global app counters
                for app in row['base_apps']:
                    self.app_counters['base'][app] += 1
                
                for app in row['variant_apps']:
                    self.app_counters['variant'][app] += 1
                
                # Count combinations
                if len(row['base_apps']) >= 2:
                    for combo in combinations(sorted(row['base_apps']), 2):
                        self.app_counters['base_combos'][combo] += 1
                
                if len(row['variant_apps']) >= 2:
                    for combo in combinations(sorted(row['variant_apps']), 2):
                        self.app_counters['variant_combos'][combo] += 1
    
    def analyze(self):
        """Main analysis function that processes file in chunks"""
        print(f"Starting chunked analysis of {self.csv_path}")
        print(f"Chunk size: {self.chunk_size} rows")
        
        rows_processed = 0
        chunk_num = 0
        
        # Get column info from first chunk
        first_chunk = pd.read_csv(self.csv_path, nrows=100)
        print(f"Columns found: {first_chunk.columns.tolist()}")
        
        # Process file in chunks
        for chunk in pd.read_csv(self.csv_path, chunksize=self.chunk_size):
            chunk_num += 1
            print(f"Processing chunk {chunk_num} ({len(chunk)} rows)...")
            
            # Process the chunk
            processed_chunk = self.process_chunk(chunk)
            
            # Update statistics
            self.update_statistics(processed_chunk)
            
            rows_processed += len(processed_chunk)
            
            # Check if we've reached the max rows
            if self.max_rows and rows_processed >= self.max_rows:
                print(f"Reached maximum rows limit: {self.max_rows}")
                break
            
            # Garbage collection to free memory
            del processed_chunk
            gc.collect()
        
        print(f"Analysis complete. Processed {rows_processed} rows in {chunk_num} chunks")
        return self.stats, self.app_counters
    
    def create_visualizations(self, output_dir='/home/claude'):
        """Create visualization files"""
        plots = []
        
        # 1. Top devices by count
        device_counts = Counter()
        for device, release_data in self.stats.items():
            for release_type, data in release_data.items():
                device_counts[device] += data['count']
        
        top_devices = device_counts.most_common(20)
        if top_devices:
            fig1 = px.bar(x=[d[1] for d in top_devices], 
                         y=[d[0] for d in top_devices],
                         orientation='h',
                         title='Top 20 Devices by Frequency',
                         labels={'x': 'Count', 'y': 'Device'})
            fig1.write_html(f'{output_dir}/device_distribution.html')
            plots.append('device_distribution.html')
        
        # 2. Release type distribution
        release_counts = Counter()
        for device, release_data in self.stats.items():
            for release_type, data in release_data.items():
                release_counts[release_type] += data['count']
        
        if release_counts:
            fig2 = px.pie(values=list(release_counts.values()), 
                         names=list(release_counts.keys()),
                         title='Release Type Distribution')
            fig2.write_html(f'{output_dir}/release_type_distribution.html')
            plots.append('release_type_distribution.html')
        
        # 3. Top base apps
        top_base = self.app_counters['base'].most_common(15)
        if top_base:
            fig3 = px.bar(x=[item[1] for item in top_base], 
                         y=[item[0] for item in top_base],
                         orientation='h',
                         title='Top 15 Apps in Base Only',
                         labels={'x': 'Frequency', 'y': 'App'})
            fig3.write_html(f'{output_dir}/top_base_apps.html')
            plots.append('top_base_apps.html')
        
        # 4. Top variant apps
        top_variant = self.app_counters['variant'].most_common(15)
        if top_variant:
            fig4 = px.bar(x=[item[1] for item in top_variant], 
                         y=[item[0] for item in top_variant],
                         orientation='h',
                         title='Top 15 Apps in Variant Only',
                         labels={'x': 'Frequency', 'y': 'App'})
            fig4.write_html(f'{output_dir}/top_variant_apps.html')
            plots.append('top_variant_apps.html')
        
        # 5. Top app combinations
        top_base_combos = self.app_counters['base_combos'].most_common(10)
        if top_base_combos:
            combo_names = [f"{c[0][0]} + {c[0][1]}" for c in top_base_combos]
            combo_counts = [c[1] for c in top_base_combos]
            fig5 = px.bar(x=combo_counts, y=combo_names,
                         orientation='h',
                         title='Top 10 App Combinations in Base',
                         labels={'x': 'Frequency', 'y': 'App Combination'})
            fig5.write_html(f'{output_dir}/top_base_combinations.html')
            plots.append('top_base_combinations.html')
        
        print(f"Created {len(plots)} visualization files")
        return plots
    
    def create_detailed_report(self, output_path='/home/claude/detailed_report.docx'):
        """Create a detailed DOCX report"""
        doc = Document()
        
        # Title
        doc.add_heading('Android Device Configuration Analysis Report', 0)
        doc.add_heading('Large-Scale Data Analysis', 1)
        
        # Summary
        doc.add_heading('Executive Summary', 1)
        
        total_devices = len(self.stats)
        total_entries = sum(sum(data['count'] for data in release_data.values()) 
                          for release_data in self.stats.values())
        
        doc.add_paragraph(f'Total unique devices analyzed: {total_devices}')
        doc.add_paragraph(f'Total configuration entries: {total_entries:,}')
        doc.add_paragraph(f'Unique apps in base: {len(self.app_counters["base"])}')
        doc.add_paragraph(f'Unique apps in variant: {len(self.app_counters["variant"])}')
        
        # Top-level statistics
        doc.add_heading('Key Findings', 1)
        
        # Most common apps
        doc.add_heading('Most Frequent Applications', 2)
        
        doc.add_heading('Top 10 Base Apps', 3)
        for app, count in self.app_counters['base'].most_common(10):
            doc.add_paragraph(f'• {app}: {count:,} occurrences', style='List Bullet')
        
        doc.add_heading('Top 10 Variant Apps', 3)
        for app, count in self.app_counters['variant'].most_common(10):
            doc.add_paragraph(f'• {app}: {count:,} occurrences', style='List Bullet')
        
        # App combinations
        doc.add_heading('Common App Combinations', 2)
        
        doc.add_heading('Top 5 Base App Pairs', 3)
        for combo, count in self.app_counters['base_combos'].most_common(5):
            doc.add_paragraph(f'• {combo[0]} + {combo[1]}: {count:,} occurrences', style='List Bullet')
        
        doc.add_heading('Top 5 Variant App Pairs', 3)
        for combo, count in self.app_counters['variant_combos'].most_common(5):
            doc.add_paragraph(f'• {combo[0]} + {combo[1]}: {count:,} occurrences', style='List Bullet')
        
        # Device-specific analysis (top 5 devices)
        doc.add_heading('Device-Specific Analysis', 1)
        
        device_totals = [(device, sum(data['count'] for data in release_data.values()))
                        for device, release_data in self.stats.items()]
        device_totals.sort(key=lambda x: x[1], reverse=True)
        
        for device, total in device_totals[:5]:
            doc.add_heading(f'Device: {device} (Total: {total:,} entries)', 2)
            
            for release_type, data in self.stats[device].items():
                doc.add_heading(f'Release Type: {release_type}', 3)
                doc.add_paragraph(f'Entries: {data["count"]:,}')
                
                # Top apps by fingerprint
                if data['fingerprint_stats']:
                    doc.add_paragraph('Top Fingerprint Configurations:', style='List Bullet')
                    
                    # Get top fingerprints by total app count
                    fp_totals = []
                    for fp, app_data in data['fingerprint_stats'].items():
                        total_apps = sum(app_data['base'].values()) + sum(app_data['variant'].values())
                        fp_totals.append((fp, total_apps, app_data))
                    
                    fp_totals.sort(key=lambda x: x[1], reverse=True)
                    
                    for fp, total, app_data in fp_totals[:3]:
                        doc.add_paragraph(f'  Fingerprint: {fp}', style='List Bullet 2')
                        
                        top_base = app_data['base'].most_common(3)
                        if top_base:
                            doc.add_paragraph(f'    Top base apps: {", ".join([f"{app} ({cnt})" for app, cnt in top_base])}', 
                                           style='List Bullet 3')
                        
                        top_variant = app_data['variant'].most_common(3)
                        if top_variant:
                            doc.add_paragraph(f'    Top variant apps: {", ".join([f"{app} ({cnt})" for app, cnt in top_variant])}', 
                                           style='List Bullet 3')
        
        # Conclusions
        doc.add_heading('Analysis Insights', 1)
        doc.add_paragraph('Based on the large-scale analysis of device configurations:')
        
        insights = []
        
        # Check for dominant apps
        top_base_app = self.app_counters['base'].most_common(1)[0] if self.app_counters['base'] else None
        if top_base_app:
            insights.append(f'The most common base app "{top_base_app[0]}" appears {top_base_app[1]:,} times')
        
        # Check for common combinations
        top_combo = self.app_counters['base_combos'].most_common(1)[0] if self.app_counters['base_combos'] else None
        if top_combo:
            insights.append(f'The app pair "{top_combo[0][0]}" and "{top_combo[0][1]}" frequently appear together ({top_combo[1]:,} times)')
        
        # Device diversity
        if len(self.stats) > 10:
            insights.append(f'High device diversity with {len(self.stats)} unique devices analyzed')
        
        for insight in insights:
            doc.add_paragraph(f'• {insight}', style='List Bullet')
        
        # Save document
        doc.save(output_path)
        print(f"Detailed report saved to {output_path}")
        
        return doc

def main():
    """Main execution with large file handling"""
    print("=" * 60)
    print("LARGE FILE ANDROID CONFIGURATION ANALYZER")
    print("=" * 60)
    
    csv_path = "output.csv"
    
    if not os.path.exists(csv_path):
        print(f"Error: {csv_path} not found")
        return
    
    # Get file size
    file_size = os.path.getsize(csv_path) / (1024 * 1024)  # Size in MB
    print(f"File size: {file_size:.2f} MB")
    
    # Determine chunk size based on file size
    if file_size > 100:
        chunk_size = 50000
    elif file_size > 50:
        chunk_size = 20000
    else:
        chunk_size = 10000
    
    # Create analyzer
    analyzer = LargeFileAnalyzer(
        csv_path=csv_path,
        chunk_size=chunk_size,
        max_rows=None  # Set to a number if you want to limit processing
    )
    
    # Run analysis
    stats, app_counters = analyzer.analyze()
    
    # Create visualizations
    print("\nCreating visualizations...")
    plots = analyzer.create_visualizations()
    
    # Create report
    print("\nCreating detailed report...")
    analyzer.create_detailed_report()
    
    print("\n" + "=" * 60)
    print("ANALYSIS COMPLETE!")
    print("=" * 60)
    print("\nGenerated files:")
    print("- detailed_report.docx : Comprehensive analysis report")
    for plot in plots:
        print(f"- {plot}")
    
    print("\nStatistics Summary:")
    print(f"- Devices analyzed: {len(stats)}")
    print(f"- Unique base apps: {len(app_counters['base'])}")
    print(f"- Unique variant apps: {len(app_counters['variant'])}")
    print(f"- Unique app combinations found: {len(app_counters['base_combos']) + len(app_counters['variant_combos'])}")

if __name__ == "__main__":
    main()
