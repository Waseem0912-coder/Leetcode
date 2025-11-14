#!/usr/bin/env python3
"""
Weekly Reports Consolidation with Chunking for Large PDFs
Handles millions of tokens by processing in manageable chunks.
"""
import argparse
from pathlib import Path
import pdfplumber
from docx import Document
from transformers import AutoTokenizer, AutoModelForCausalLM
import torch
import re
from datetime import datetime
from typing import List, Tuple

def print_environment_info():
    """Display GPU configuration."""
    print("=" * 70)
    print("🔍 GPU CONFIGURATION")
    print("=" * 70)
    
    if torch.cuda.is_available():
        num_gpus = torch.cuda.device_count()
        print(f"\nDetected {num_gpus} GPUs:")
        total_memory = 0
        for i in range(num_gpus):
            gpu_mem = torch.cuda.get_device_properties(i).total_memory / 1e9
            total_memory += gpu_mem
            print(f"  GPU {i}: {torch.cuda.get_device_name(i)} - {gpu_mem:.1f} GB")
        print(f"\nTotal VRAM: {total_memory:.1f} GB")
    else:
        print("\n⚠️  No GPU detected - will run on CPU")
    print("=" * 70 + "\n")


def load_model_and_tokenizer(model_id: str = "openai/gpt-oss-120b"):
    """Load model and tokenizer distributed across all GPUs."""
    print(f"🔧 Loading {model_id}...")
    
    num_gpus = torch.cuda.device_count()
    if num_gpus == 0:
        raise RuntimeError("No CUDA GPUs available")
    
    # Configure memory per GPU
    max_memory_per_gpu = {}
    for i in range(num_gpus):
        total_mem = torch.cuda.get_device_properties(i).total_memory / 1e9
        usable_mem = max(total_mem * 0.5, 60)  # Use 50% for model, leave rest for operations
        max_memory_per_gpu[i] = f"{int(usable_mem)}GiB"
    
    print(f"\n💾 Memory allocation per GPU: {list(max_memory_per_gpu.values())[0]}")
    
    # Load tokenizer
    print("📥 Loading tokenizer...")
    tokenizer = AutoTokenizer.from_pretrained(model_id, trust_remote_code=True)
    
    # Load model
    print("📥 Loading model...")
    model = AutoModelForCausalLM.from_pretrained(
        model_id,
        torch_dtype=torch.bfloat16,
        device_map="balanced",
        max_memory=max_memory_per_gpu,
        trust_remote_code=True,
        low_cpu_mem_usage=True,
    )
    
    # Verify distribution
    if hasattr(model, 'hf_device_map'):
        print("\n📍 Model Distribution:")
        device_counts = {}
        for name, device in model.hf_device_map.items():
            device_str = f"GPU {device}" if isinstance(device, int) else str(device)
            device_counts[device_str] = device_counts.get(device_str, 0) + 1
        
        for device, count in sorted(device_counts.items()):
            print(f"  {device}: {count} layers")
    
    print("\n✅ Model and tokenizer loaded\n")
    return model, tokenizer


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract all text from PDF."""
    print(f"📄 Extracting text from: {pdf_path.name}")
    
    all_text = []
    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        print(f"   Found {total_pages} pages")
        
        for i, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text()
            if page_text:
                all_text.append(f"\n{'='*70}\nPAGE {i}\n{'='*70}\n")
                all_text.append(page_text.strip())
            
            if i % 10 == 0 or i == total_pages:
                print(f"   Progress: {i}/{total_pages} pages...", end='\r')
        
        print()
    
    combined_text = "\n\n".join(all_text)
    print(f"✅ Extracted {total_pages} pages ({len(combined_text):,} characters)\n")
    return combined_text


def count_tokens(text: str, tokenizer) -> int:
    """Count tokens in text."""
    tokens = tokenizer.encode(text, add_special_tokens=True)
    return len(tokens)


def chunk_text_by_tokens(text: str, tokenizer, max_input_tokens: int = 90000) -> List[str]:
    """
    Split text into chunks based on token count.
    Tries to split on week boundaries or page markers for clean breaks.
    """
    print(f"\n📦 Chunking text (max {max_input_tokens:,} tokens per chunk)...")
    
    # First, try to split by weeks
    week_pattern = r'\n={70}\nPAGE \d+\n={70}\n'
    sections = re.split(week_pattern, text)
    
    chunks = []
    current_chunk = []
    current_tokens = 0
    
    for i, section in enumerate(sections):
        if not section.strip():
            continue
            
        section_tokens = count_tokens(section, tokenizer)
        
        # If single section is too large, split it further
        if section_tokens > max_input_tokens:
            print(f"   ⚠️  Section {i} has {section_tokens:,} tokens (too large)")
            # Split by paragraphs
            paragraphs = section.split('\n\n')
            for para in paragraphs:
                para_tokens = count_tokens(para, tokenizer)
                if current_tokens + para_tokens > max_input_tokens and current_chunk:
                    # Save current chunk
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = [para]
                    current_tokens = para_tokens
                else:
                    current_chunk.append(para)
                    current_tokens += para_tokens
        else:
            # Check if adding this section exceeds limit
            if current_tokens + section_tokens > max_input_tokens and current_chunk:
                # Save current chunk
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = [section]
                current_tokens = section_tokens
            else:
                current_chunk.append(section)
                current_tokens += section_tokens
    
    # Add remaining chunk
    if current_chunk:
        chunks.append('\n\n'.join(current_chunk))
    
    # Verify chunk sizes
    print(f"\n   Created {len(chunks)} chunks:")
    for i, chunk in enumerate(chunks, 1):
        chunk_tokens = count_tokens(chunk, tokenizer)
        print(f"   Chunk {i}: {chunk_tokens:,} tokens ({len(chunk):,} chars)")
    
    return chunks


def generate_text(model, tokenizer, prompt: str, max_new_tokens: int = 5000) -> str:
    """
    Generate text using ALL GPUs properly.
    """
    
    # Tokenize with truncation
    inputs = tokenizer(
        prompt, 
        return_tensors="pt", 
        truncation=True, 
        max_length=120000  # Model's context window
    )
    
    # Get first device dynamically
    first_device = None
    if hasattr(model, 'hf_device_map'):
        for name, device in model.hf_device_map.items():
            first_device = device
            break
    
    if first_device is None:
        first_device = next(model.parameters()).device
    
    inputs = {k: v.to(first_device) for k, v in inputs.items()}
    
    input_tokens = inputs['input_ids'].shape[1]
    
    # Check if input fits
    if input_tokens >= 120000:
        print(f"   ⚠️  WARNING: Input ({input_tokens:,} tokens) at context limit!")
    
    print(f"   Input: {input_tokens:,} tokens")
    print(f"   Max output: {max_new_tokens:,} tokens")
    print(f"   Processing on device: {first_device}")
    
    # Generate
    with torch.inference_mode():
        outputs = model.generate(
            **inputs,
            max_new_tokens=max_new_tokens,
            do_sample=False,
            temperature=1.0,
            pad_token_id=tokenizer.eos_token_id,
            eos_token_id=tokenizer.eos_token_id,
        )
    
    # Decode only new tokens
    generated_text = tokenizer.decode(
        outputs[0][inputs['input_ids'].shape[1]:], 
        skip_special_tokens=True
    )
    
    return generated_text


def process_chunk(chunk: str, chunk_num: int, total_chunks: int, model, tokenizer, max_new_tokens: int = 5000) -> str:
    """Process a single chunk with your exact prompt."""
    
    print(f"\n🔄 Processing chunk {chunk_num}/{total_chunks}...")
    
    prompt = f"""You are analyzing a long text containing ~45 weeks of weekly reports.  
Each week contains several main bullets (topics/projects) and sub-bullets (details, updates, tasks).  
Some topics appear in multiple weeks with different updates, some appear only once, and some weeks reset or replace previous information.

GOAL  
Reconstruct the entire set of weekly reports into a clean, structured markdown document that:
- Preserves the original week-by-week structure
- Preserves the original topic names EXACTLY as they appeared (no normalization or merging)
- Preserves all details under each week and under the correct topic
- Makes the final document readable, scannable, and consistently formatted

INTERNAL REASONING STEPS (DO NOT OUTPUT THESE STEPS)  
1. Detect week boundaries (Week 1 → Week 45 or similar markers).  
2. Inside each week:
   - Identify which bullets are main topics (first-level bullets or bolded headers).
   - Identify which bullets are subpoints or child items.
3. Do NOT normalize, merge, rename, or cluster topics.
   - If Week 3 says "Offline Monitoring" and Week 7 says "Monitoring Improvements," treat them as two separate topics even if they seem similar.
4. Do NOT build a cross-week topic timeline.
   - Each week is self-contained.
   - Simply reconstruct Week 1's content, then Week 2's content, etc.
5. Preserve:
   - All numbers, dates, metrics, names, file paths, schemas, decisions, blockers, achievements.
   - All bullet points and nested bullets.
   - The order bullets appeared in the original text.
6. Clean up:
   - Remove PDF artifacts, page numbers, broken lines, nonsense breaks.
   - Convert inconsistent indentation into consistent markdown hierarchy.

OUTPUT FORMAT  
Use the following exact structure:
- `# Week X` for each week  
- Under each week, use `## <topic name exactly as written>`  
- Under each topic, use bullet points (`-`) and nested bullets as needed  
- Keep chronological order (Week 1 → Week 45)  
- Keep the original content attached to its week only  
- Do NOT merge or reorganize topics across weeks  
- Do NOT infer missing continuity or relationships

ADDITIONAL REQUIREMENTS  
- Output ONLY the final markdown.  
- NO explanations, NO preamble, NO reasoning.
- No hallucination: use only the information in the provided text.

TEXT TO CONVERT (CHUNK {chunk_num}/{total_chunks}):
{chunk}

OUTPUT:
"""
    
    result = generate_text(model, tokenizer, prompt, max_new_tokens)
    print(f"✅ Chunk {chunk_num} processed ({len(result):,} characters)")
    
    return result


def combine_chunks(chunk_results: List[str]) -> str:
    """Combine processed chunks into final document."""
    print(f"\n🔗 Combining {len(chunk_results)} chunks...")
    
    # Simply concatenate - each chunk should be self-contained weeks
    combined = "\n\n".join(chunk_results)
    
    print(f"✅ Combined into {len(combined):,} characters\n")
    return combined


def markdown_to_docx(markdown_text: str, output_path: Path):
    """Convert markdown to Word document."""
    print(f"📝 Creating Word document...")
    
    doc = Document()
    
    # Title page
    doc.add_heading('Weekly Reports - Reconstructed', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_page_break()
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
        
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 4)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            # Count indentation
            indent_level = 0
            temp_line = line
            while temp_line.startswith('  '):
                indent_level += 1
                temp_line = temp_line[2:]
            
            content = line.lstrip('- *').strip()
            
            # Handle bold
            if '**' in content:
                p = doc.add_paragraph(style='List Bullet' if indent_level == 0 else 'List Bullet 2')
                parts = content.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)
                    else:
                        p.add_run(part).bold = True
            else:
                style = 'List Bullet' if indent_level == 0 else 'List Bullet 2'
                doc.add_paragraph(content, style=style)
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            content = re.sub(r'^\d+\.\s', '', line)
            doc.add_paragraph(content, style='List Number')
        
        # Regular paragraphs
        else:
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)
                    else:
                        p.add_run(part).bold = True
            else:
                doc.add_paragraph(line)
    
    doc.save(str(output_path))
    print(f"✅ Saved: {output_path}\n")


def main():
    """Main workflow with chunking support."""
    
    parser = argparse.ArgumentParser(
        description='Consolidate weekly reports with automatic chunking for large PDFs'
    )
    parser.add_argument('--pdf', type=str, required=True, help='Path to PDF')
    parser.add_argument('--output', type=str, default='./output', help='Output directory')
    parser.add_argument('--model', type=str, default='openai/gpt-oss-120b', help='Model ID')
    parser.add_argument('--max-output-tokens', type=int, default=5000, help='Max output tokens per chunk')
    parser.add_argument('--max-input-tokens', type=int, default=90000, help='Max input tokens per chunk')
    
    args = parser.parse_args()
    
    pdf_path = Path(args.pdf)
    output_dir = Path(args.output)
    
    if not pdf_path.exists():
        print(f"❌ ERROR: PDF not found: {pdf_path}")
        return 1
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print_environment_info()
    
    print("=" * 70)
    print("📊 WEEKLY REPORTS RECONSTRUCTION WITH CHUNKING")
    print("=" * 70)
    print(f"\n📄 Input: {pdf_path}")
    print(f"📁 Output: {output_dir}")
    print(f"⚙️  Max input tokens: {args.max_input_tokens:,}")
    print(f"⚙️  Max output tokens per chunk: {args.max_output_tokens:,}\n")
    
    # Load model and tokenizer
    model, tokenizer = load_model_and_tokenizer(args.model)
    
    # Extract PDF
    raw_text = extract_text_from_pdf(pdf_path)
    
    # Save raw text
    raw_path = output_dir / "01_raw_extracted_text.txt"
    raw_path.write_text(raw_text, encoding='utf-8')
    print(f"💾 Saved: {raw_path}\n")
    
    # Count tokens in full text
    print("🔢 Counting tokens in full text...")
    total_tokens = count_tokens(raw_text, tokenizer)
    print(f"   Total tokens: {total_tokens:,}")
    print(f"   Total characters: {len(raw_text):,}")
    print(f"   Total words: {len(raw_text.split()):,}\n")
    
    # Check if we need chunking
    if total_tokens > args.max_input_tokens:
        print(f"⚠️  Input ({total_tokens:,} tokens) exceeds limit ({args.max_input_tokens:,} tokens)")
        print(f"   Will process in chunks...\n")
        
        # Chunk the text
        chunks = chunk_text_by_tokens(raw_text, tokenizer, args.max_input_tokens)
        
        # Process each chunk
        chunk_results = []
        for i, chunk in enumerate(chunks, 1):
            result = process_chunk(chunk, i, len(chunks), model, tokenizer, args.max_output_tokens)
            chunk_results.append(result)
            
            # Save intermediate result
            chunk_path = output_dir / f"chunk_{i:02d}_result.md"
            chunk_path.write_text(result, encoding='utf-8')
            print(f"   💾 Saved intermediate: {chunk_path}")
        
        # Combine chunks
        final_markdown = combine_chunks(chunk_results)
        
    else:
        print(f"✅ Input fits in single chunk ({total_tokens:,} tokens)\n")
        
        # Process in single chunk
        final_markdown = process_chunk(raw_text, 1, 1, model, tokenizer, args.max_output_tokens)
    
    # Save final markdown
    md_path = output_dir / "02_reconstructed_report.md"
    md_path.write_text(final_markdown, encoding='utf-8')
    print(f"💾 Saved: {md_path}\n")
    
    # Create Word document
    docx_path = output_dir / "03_reconstructed_report.docx"
    markdown_to_docx(final_markdown, docx_path)
    
    print("=" * 70)
    print("✅ COMPLETE")
    print("=" * 70)
    print(f"\n📁 Outputs in {output_dir}:")
    print(f"   1. {raw_path.name} - Raw extracted text")
    print(f"   2. {md_path.name} - Reconstructed markdown")
    print(f"   3. {docx_path.name} - Word document")
    
    if total_tokens > args.max_input_tokens:
        print(f"\n   📦 Also saved {len(chunks)} intermediate chunk results")
    
    print(f"\n📊 Statistics:")
    print(f"   Input: {total_tokens:,} tokens")
    print(f"   Output: {len(final_markdown):,} characters")
    print(f"\n🎉 Your weekly reports have been reconstructed!\n")
    
    return 0


if __name__ == "__main__":
    exit(main())
